#!/usr/bin/env python3
"""Quick DOCX inspection helper for the builtin doc skill.

Prints paragraph counts, sample paragraphs, and table summaries.
Requires `python-docx`.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect a DOCX file with python-docx.")
    parser.add_argument("docx_path", help="Path to the .docx file")
    parser.add_argument(
        "--paragraphs",
        type=int,
        default=8,
        help="Number of paragraphs to sample",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    if not docx_path.is_file():
        print(f"error: file not found: {docx_path}", file=sys.stderr)
        return 1

    if docx_path.suffix.lower() != ".docx":
        print(
            "error: inspect_docx.py expects a .docx file. Convert legacy .doc first.",
            file=sys.stderr,
        )
        return 2

    try:
        from docx import Document
    except ImportError:
        print(
            "error: python-docx is not installed. Install with `uv pip install python-docx` "
            "or `python3 -m pip install python-docx`.",
            file=sys.stderr,
        )
        return 3

    try:
        document = Document(str(docx_path))
    except Exception as exc:  # pragma: no cover - defensive runtime wrapper
        print(f"error: failed to open DOCX: {exc}", file=sys.stderr)
        return 4

    print(f"path: {docx_path}")
    print(f"paragraphs: {len(document.paragraphs)}")
    print(f"tables: {len(document.tables)}")

    print("sampled_paragraphs:")
    sampled = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        sampled += 1
        print(f"  {sampled}: {text}")
        if sampled >= args.paragraphs:
            break

    if document.tables:
        print("table_summaries:")
        for index, table in enumerate(document.tables, start=1):
            rows = len(table.rows)
            cols = len(table.columns)
            print(f"  - table {index}: rows={rows} cols={cols}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
